from __future__ import annotations

import io
import os


ALLOWED_EXTENSIONS = frozenset({
    ".txt", ".md", ".markdown", ".text",
    ".docx", ".xlsx", ".pdf",
})


def is_allowed(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(content: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext in (".txt", ".md", ".markdown", ".text"):
            return _extract_txt(content)
        if ext == ".docx":
            return _extract_docx(content)
        if ext == ".xlsx":
            return _extract_xlsx(content)
        if ext == ".pdf":
            return _extract_pdf(content)
    except Exception as exc:
        raise ValueError(f"无法解析文件 {filename}: {exc}") from exc
    raise ValueError(f"不支持的文件类型: {ext}")


def _extract_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _extract_docx(content: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = "\t".join(cells)
            if line.strip():
                parts.append(line)

    return "\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []

    for ws in wb.worksheets:
        parts.append(f"[{ws.title}]")
        for row in ws.iter_rows():
            cells = []
            for cell in row:
                v = cell.value
                cells.append(str(v) if v is not None else "")
            line = "\t".join(cells)
            if line.strip():
                parts.append(line)

    wb.close()
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    import pdfplumber

    parts: list[str] = []

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)

            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    cells = [cell.strip() if cell else "" for cell in row]
                    line = "\t".join(cells)
                    if line.strip():
                        parts.append(line)

    return "\n".join(parts)
